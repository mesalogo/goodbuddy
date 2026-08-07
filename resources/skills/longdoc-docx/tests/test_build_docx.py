import importlib.util
import json
import tempfile
import unittest
from pathlib import Path

from docx import Document
from PIL import Image


SKILL_DIR = Path(__file__).resolve().parents[1]
SCRIPT_PATH = SKILL_DIR / "scripts" / "build_docx.py"
SPEC = importlib.util.spec_from_file_location("build_docx", SCRIPT_PATH)
build_docx = importlib.util.module_from_spec(SPEC)
SPEC.loader.exec_module(build_docx)


class BuildDocxTests(unittest.TestCase):
    def write_config(self, root, config):
        path = root / "document.json"
        path.write_text(json.dumps(config, ensure_ascii=False), encoding="utf-8")
        return path

    def test_builds_supported_markdown_elements(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            image_path = root / "diagram.png"
            Image.new("RGB", (400, 200), "white").save(image_path, dpi=(200, 200))
            chapter = root / "chapter.md"
            chapter.write_text(
                "# 一、概述\n\n"
                "正文包含**粗体**、*斜体*和`代码`。\n\n"
                "- 列表一\n- 列表二\n\n"
                "| 项目 | 详细说明 |\n|---|---|\n| A | 一段较长的内容 |\n\n"
                "![架构图](diagram.png)\n\n"
                "*图 1 架构图*\n\n"
                "```python\nprint('ok')\n```\n\n"
                "> 引用说明\n",
                encoding="utf-8",
            )
            config_path = self.write_config(
                root,
                {
                    "title": "测试文档",
                    "subtitle": "构建验证",
                    "output": "out/test.docx",
                    "toc_depth": 0,
                    "chapters": [{"path": "chapter.md"}],
                },
            )
            loaded_path, config = build_docx.load_config(config_path)
            output = build_docx.build(loaded_path, config)

            self.assertTrue(output.is_file())
            document = Document(output)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertIn("一、概述", text)
            self.assertIn("图 1 架构图", text)
            self.assertIn("print('ok')", text)
            self.assertEqual(len(document.tables), 1)
            images = [
                rel
                for rel in document.part.rels.values()
                if "image" in rel.reltype
            ]
            self.assertEqual(len(images), 1)

    def test_landscape_widens_tables(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            header = "| " + " | ".join(f"列{i}" for i in range(11)) + " |"
            divider = "|" + "---|" * 11
            row = "| " + " | ".join(f"值{i}" for i in range(11)) + " |"
            (root / "wide.md").write_text(
                f"# 宽表\n\n{header}\n{divider}\n{row}\n", encoding="utf-8"
            )
            widths = {}
            for mode in ("portrait", "landscape"):
                config_path = self.write_config(
                    root,
                    {
                        "title": "宽表测试",
                        "output": f"out/{mode}.docx",
                        "toc_depth": 0,
                        "orientation": mode,
                        "chapters": [{"path": "wide.md"}],
                    },
                )
                loaded, config = build_docx.load_config(config_path)
                document = Document(build_docx.build(loaded, config))
                widths[mode] = sum(
                    cell.width.cm for cell in document.tables[0].rows[0].cells
                )
            self.assertGreater(widths["landscape"], widths["portrait"] + 5)

    def test_wide_table_fits_longest_word_in_every_column(self):
        from bs4 import BeautifulSoup

        headers = [
            "需求编号", "需求出处", "需求原文", "类别", "强制/评分", "响应状态",
            "实现说明", "偏离说明", "证据编号", "方案章节", "验证方法",
        ]
        body = [
            "REQ-001", "示例技术要求 3.1.1",
            "系统应提供与 OpenAI 接口兼容的统一调用入口。", "功能", "强制",
            "compliant", "网关提供模型列表、对话补全和向量化四类接口，统一鉴权",
            "无", "FEAT-001", "三.1、四.1", "依次调用四类接口并核对返回结构",
        ]
        head = "".join(f"<th>{h}</th>" for h in headers)
        cells = "".join(f"<td>{c}</td>" for c in body)
        rows = BeautifulSoup(
            f"<table><tr>{head}</tr><tr>{cells}</tr></table>", "html.parser"
        ).find_all("tr")
        widths = build_docx.compute_col_widths(rows, 11, 25.7, 9.5)

        self.assertAlmostEqual(sum(widths), 25.7, places=3)
        em_cm = 9.5 / 28.35
        for index, (header, cell) in enumerate(zip(headers, body)):
            longest = max(
                build_docx.em_width(word)
                for text in (header, cell)
                for word in text.split()
            )
            self.assertGreaterEqual(
                widths[index] + 1e-6,
                min(em_cm * longest, 25.7 / 11),
                f"column {index} ({header}) truncates its longest word",
            )
        self.assertGreater(widths[6], widths[3])

    def test_blockquote_keeps_paragraph_breaks(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            (root / "quote.md").write_text(
                "# 引用\n\n> 第一段说明。\n>\n> 第二段说明。\n", encoding="utf-8"
            )
            config_path = self.write_config(
                root,
                {
                    "title": "引用测试",
                    "output": "out/quote.docx",
                    "toc_depth": 0,
                    "chapters": [{"path": "quote.md"}],
                },
            )
            loaded, config = build_docx.load_config(config_path)
            document = Document(build_docx.build(loaded, config))
            texts = [p.text for p in document.paragraphs]
            self.assertIn("第一段说明。", texts)
            self.assertIn("第二段说明。", texts)

    def test_only_forced_chapter_boundary_adds_page_break(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            (root / "one.md").write_text("# 第一章\n", encoding="utf-8")
            (root / "two.md").write_text("# 第二章\n", encoding="utf-8")
            config_path = self.write_config(
                root,
                {
                    "title": "分页测试",
                    "cover": False,
                    "toc_depth": 0,
                    "output": "test.docx",
                    "chapters": [
                        {"path": "one.md", "page_break_before": False},
                        {"path": "two.md", "page_break_before": True},
                    ],
                },
            )
            loaded_path, config = build_docx.load_config(config_path)
            output = build_docx.build(loaded_path, config)
            document = Document(output)
            self.assertEqual(document._element.xml.count('w:type="page"'), 1)

    def test_rejects_missing_chapters(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            config_path = self.write_config(
                root,
                {"title": "无章节", "output": "test.docx", "chapters": []},
            )
            with self.assertRaisesRegex(ValueError, "chapters"):
                build_docx.load_config(config_path)

    def test_rejects_non_boolean_page_break(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            (root / "chapter.md").write_text("# 章节\n", encoding="utf-8")
            config_path = self.write_config(
                root,
                {
                    "title": "错误分页配置",
                    "output": "test.docx",
                    "chapters": [
                        {"path": "chapter.md", "page_break_before": "false"}
                    ],
                },
            )
            loaded_path, config = build_docx.load_config(config_path)
            with self.assertRaisesRegex(ValueError, "page_break_before"):
                build_docx.build(loaded_path, config)

    def test_rejects_mixed_text_and_image_paragraph(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            Image.new("RGB", (100, 50), "white").save(root / "diagram.png")
            (root / "chapter.md").write_text(
                "说明文字 ![架构图](diagram.png)\n",
                encoding="utf-8",
            )
            config_path = self.write_config(
                root,
                {
                    "title": "图片格式测试",
                    "cover": False,
                    "toc_depth": 0,
                    "output": "test.docx",
                    "chapters": ["chapter.md"],
                },
            )
            loaded_path, config = build_docx.load_config(config_path)
            with self.assertRaisesRegex(ValueError, "图片必须独占"):
                build_docx.build(loaded_path, config)


if __name__ == "__main__":
    unittest.main()
