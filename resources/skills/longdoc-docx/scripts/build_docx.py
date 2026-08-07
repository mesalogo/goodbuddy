#!/usr/bin/env python3
"""Build a styled DOCX from an explicit Markdown chapter manifest."""

import argparse
import json
import re
import unicodedata
from pathlib import Path

import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image


DEFAULTS = {
    "body_font_zh": "宋体",
    "body_font_en": "Times New Roman",
    "heading_font_zh": "黑体",
    "heading_color": "1F3864",
    "body_size": 11.5,
    "toc_depth": 3,
    "max_image_width_cm": 14.66,
}
HEADING_SIZES = {1: 18, 2: 15, 3: 13, 4: 12, 5: 11.5, 6: 11.5}


def parse_args():
    parser = argparse.ArgumentParser(description="将多章节 Markdown 构建为 DOCX")
    parser.add_argument("--config", required=True, help="document.json 路径")
    return parser.parse_args()


def load_config(path):
    config_path = Path(path).expanduser().resolve()
    with config_path.open(encoding="utf-8") as handle:
        config = json.load(handle)
    if not isinstance(config, dict):
        raise ValueError("配置根节点必须是 JSON 对象")
    for key in ("title", "output"):
        if not isinstance(config.get(key), str) or not config[key].strip():
            raise ValueError(f"{key} 必须是非空字符串")
    if not isinstance(config.get("chapters"), list) or not config["chapters"]:
        raise ValueError("chapters 必须是非空数组")
    merged = {**DEFAULTS, **config}
    if not isinstance(merged["toc_depth"], int) or isinstance(merged["toc_depth"], bool):
        raise ValueError("toc_depth 必须是整数")
    for key in ("body_size", "max_image_width_cm"):
        value = merged[key]
        if not isinstance(value, (int, float)) or isinstance(value, bool) or value <= 0:
            raise ValueError(f"{key} 必须是正数")
    return config_path, merged


def rgb(value):
    value = value.lstrip("#")
    if not re.fullmatch(r"[0-9a-fA-F]{6}", value):
        raise ValueError(f"颜色必须是六位十六进制值：{value!r}")
    return RGBColor.from_string(value.upper())


def add_font(
    run,
    config,
    size=None,
    bold=False,
    italic=False,
    color=None,
    code=False,
    en_font=None,
    zh_font=None,
):
    en_font = en_font or ("Consolas" if code else config["body_font_en"])
    zh_font = zh_font or ("Consolas" if code else config["body_font_zh"])
    run.font.name = en_font
    run.font.size = Pt(size or config["body_size"])
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), en_font)
    rfonts.set(qn("w:hAnsi"), en_font)
    rfonts.set(qn("w:eastAsia"), zh_font)


def add_field(paragraph, instruction, placeholder=None):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    if placeholder:
        text = OxmlElement("w:t")
        text.text = placeholder
        run._r.append(text)
    run._r.append(end)


def add_shading(target, fill):
    properties = (
        target._tc.get_or_add_tcPr()
        if hasattr(target, "_tc")
        else target._p.get_or_add_pPr()
    )
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_table_borders(table):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "B0B0B0")
        borders.append(element)
    table._tbl.tblPr.append(borders)


def setup_document(doc, config):
    section = doc.sections[0]
    if config.get("orientation", "portrait") == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21)
    else:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
    section.top_margin = Cm(config.get("margin_top_cm", 2.54))
    section.bottom_margin = Cm(config.get("margin_bottom_cm", 2.54))
    section.left_margin = Cm(config.get("margin_left_cm", 3.17))
    section.right_margin = Cm(config.get("margin_right_cm", 3.17))

    normal = doc.styles["Normal"]
    normal.font.name = config["body_font_en"]
    normal.font.size = Pt(config["body_size"])
    normal.paragraph_format.line_spacing = config.get("line_spacing", 1.4)
    normal.paragraph_format.space_after = Pt(8)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), config["body_font_zh"])

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(footer, "PAGE")


def add_cover(doc, config):
    if config.get("cover", True) is False:
        return
    landscape = config.get("orientation", "portrait") == "landscape"
    for _ in range(config.get("cover_top_spacers", 3 if landscape else 6)):
        doc.add_paragraph()
    for text, size in (
        (config["title"], 26),
        (config.get("subtitle", ""), 22),
    ):
        if not text:
            continue
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        add_font(
            run,
            config,
            size=size,
            bold=True,
            color=rgb(config["heading_color"]),
            zh_font=config["heading_font_zh"],
        )
    for _ in range(config.get("cover_middle_spacers", 4 if landscape else 8)):
        doc.add_paragraph()
    for field in ("author", "date"):
        text = config.get(field, "")
        if text:
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_font(paragraph.add_run(text), config, size=14)
    doc.add_page_break()


def add_toc(doc, config):
    depth = int(config.get("toc_depth", 3))
    if depth <= 0:
        return
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(config.get("toc_title", "目  录"))
    add_font(
        run,
        config,
        size=18,
        bold=True,
        color=rgb(config["heading_color"]),
        zh_font=config["heading_font_zh"],
    )
    doc.add_paragraph()
    paragraph = doc.add_paragraph()
    add_field(
        paragraph,
        f'TOC \\o "1-{depth}" \\h \\z \\u',
        "右键点击此处选择“更新域”以生成目录",
    )
    doc.add_page_break()


def add_inline_runs(paragraph, node, config, bold=False, italic=False):
    for child in node.children:
        name = getattr(child, "name", None)
        if name is None:
            text = str(child).replace("\n", "")
            if text:
                add_font(
                    paragraph.add_run(text),
                    config,
                    bold=bold,
                    italic=italic,
                )
        elif name in ("strong", "b"):
            add_inline_runs(paragraph, child, config, bold=True, italic=italic)
        elif name in ("em", "i"):
            add_inline_runs(paragraph, child, config, bold=bold, italic=True)
        elif name == "code":
            run = paragraph.add_run(child.get_text())
            add_font(
                run,
                config,
                size=config["body_size"] - 0.5,
                bold=bold,
                italic=italic,
                color=RGBColor(0xA0, 0x30, 0x30),
                code=True,
            )
        elif name == "br":
            paragraph.add_run().add_break()
        else:
            add_inline_runs(paragraph, child, config, bold=bold, italic=italic)


def add_heading(doc, level, text, config):
    paragraph = doc.add_paragraph(style=f"Heading {min(level, 9)}")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    paragraph.paragraph_format.space_after = Pt(8 if level == 1 else 6)
    run = paragraph.add_run(text)
    add_font(
        run,
        config,
        size=HEADING_SIZES.get(level, 11.5),
        bold=True,
        color=rgb(config["heading_color"]) if level <= 2 else RGBColor(0, 0, 0),
        zh_font=config["heading_font_zh"],
    )


def add_paragraph(doc, node, config):
    paragraph = doc.add_paragraph()
    add_inline_runs(paragraph, node, config)
    paragraph.paragraph_format.line_spacing = config.get("line_spacing", 1.4)
    paragraph.paragraph_format.space_after = Pt(8)


def add_list(doc, node, config, level=0):
    ordered = node.name == "ol"
    style = "List Number" if ordered else "List Bullet"
    for item in node.find_all("li", recursive=False):
        paragraph = doc.add_paragraph(style=style)
        paragraph.paragraph_format.left_indent = Cm(0.5 + level * 0.6)
        paragraph.paragraph_format.space_after = Pt(4)
        for child in item.children:
            if getattr(child, "name", None) in ("ul", "ol"):
                continue
            if getattr(child, "name", None) is None:
                text = str(child).replace("\n", "")
                if text:
                    add_font(paragraph.add_run(text), config)
            else:
                add_inline_runs(paragraph, child, config)
        for nested in item.find_all(["ul", "ol"], recursive=False):
            add_list(doc, nested, config, level + 1)


def display_width(text):
    return sum(2 if unicodedata.east_asian_width(char) in ("W", "F") else 1 for char in text)


def em_width(text):
    total = 0.0
    for char in text:
        if unicodedata.east_asian_width(char) in ("W", "F"):
            total += 1.0
        elif char.isupper() or char.isdigit():
            total += 0.62
        else:
            total += 0.5
    return total


def compute_col_widths(rows, ncols, content_width_cm, body_size_pt=10.5):
    lengths = [1] * ncols
    longest_word = [1] * ncols
    for row in rows:
        for index, cell in enumerate(row.find_all(["th", "td"], recursive=False)):
            if index < ncols:
                text = cell.get_text(" ", strip=True)
                lengths[index] = max(lengths[index], min(display_width(text), 160))
                longest_word[index] = max(
                    longest_word[index],
                    max((em_width(word) for word in text.split()), default=1.0),
                )
    em_cm = body_size_pt / 28.35
    padding_cm = 0.4
    floors = [
        min(em_cm * word + padding_cm, content_width_cm / ncols)
        for word in longest_word
    ]
    maximum = max(max(floors), content_width_cm * 0.55)
    widths = [None] * ncols
    remaining = content_width_cm
    pending = set(range(ncols))
    while pending:
        weight = sum(lengths[i] for i in pending)
        clamped = False
        for index in sorted(pending):
            share = remaining * lengths[index] / weight
            floor = floors[index]
            bound = floor if share < floor else (maximum if share > maximum else None)
            if bound is not None:
                widths[index] = bound
                remaining -= bound
                pending.discard(index)
                clamped = True
                break
        if not clamped:
            for index in pending:
                widths[index] = remaining * lengths[index] / weight
            break
    total = sum(widths)
    if total > content_width_cm:
        widths = [width * content_width_cm / total for width in widths]
    return widths


def set_col_widths(table, widths):
    table.autofit = False
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    else:
        for child in list(grid):
            grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(int(Cm(width).twips)))
        grid.append(column)
    for row in table.rows:
        cells = row.cells
        for index, width in enumerate(widths):
            if index < len(cells):
                cells[index].width = Cm(width)


def add_table(doc, node, config):
    rows = node.find_all("tr")
    if not rows:
        return
    ncols = max(len(row.find_all(["th", "td"], recursive=False)) for row in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    section = doc.sections[-1]
    content_width = (
        section.page_width.cm - section.left_margin.cm - section.right_margin.cm
    )
    table_size = float(config.get("table_size", config["body_size"]))
    widths = compute_col_widths(rows, ncols, content_width, table_size)
    set_table_borders(table)
    for row_index, (row_node, table_row) in enumerate(zip(rows, table.rows)):
        cell_nodes = row_node.find_all(["th", "td"], recursive=False)
        table_cells = table_row.cells
        for column_index, cell_node in enumerate(cell_nodes):
            if column_index >= len(table_cells):
                break
            cell = table_cells[column_index]
            cell.text = ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            is_header = row_index == 0
            add_inline_runs(paragraph, cell_node, config, bold=is_header)
            for run in paragraph.runs:
                add_font(
                    run,
                    config,
                    size=table_size,
                    bold=is_header or bool(run.font.bold),
                    italic=bool(run.font.italic),
                    color=RGBColor(255, 255, 255) if is_header else None,
                )
            if is_header:
                add_shading(cell, config["heading_color"].lstrip("#"))
            elif row_index % 2 == 0:
                add_shading(cell, "F2F2F2")
    set_col_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_code_block(doc, text, config):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.5)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(10)
    lines = text.rstrip("\n").split("\n")
    for index, line in enumerate(lines):
        run = paragraph.add_run(line or " ")
        add_font(
            run,
            config,
            size=9.5,
            color=RGBColor(0x33, 0x33, 0x33),
            code=True,
        )
        if index < len(lines) - 1:
            run.add_break()
    add_shading(paragraph, "F5F5F5")


def add_blockquote(doc, node, config):
    blocks = [
        child.get_text(" ", strip=True)
        for child in node.find_all("p", recursive=False)
    ]
    if not blocks:
        blocks = [node.get_text(" ", strip=True)]
    blocks = [text for text in blocks if text]
    for index, text in enumerate(blocks):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.8)
        paragraph.paragraph_format.space_after = Pt(
            10 if index == len(blocks) - 1 else 4
        )
        run = paragraph.add_run(text)
        add_font(
            run, config, size=10.5, italic=True, color=RGBColor(0x40, 0x40, 0x40)
        )
        add_shading(paragraph, "F7F7F7")


def add_image(doc, src, base_dir, config):
    image_path = (base_dir / src).resolve()
    if not image_path.is_file():
        raise FileNotFoundError(f"图片不存在：{image_path}")
    with Image.open(image_path) as image:
        width_px = image.width
        dpi = image.info.get("dpi", (150, 150))[0] or 150
    natural_width_cm = width_px / dpi * 2.54
    width_cm = min(natural_width_cm, float(config["max_image_width_cm"]))
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Cm(width_cm))
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)


def add_caption(doc, text, config):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    add_font(run, config, size=10, italic=True, color=RGBColor(0x40, 0x40, 0x40))
    paragraph.paragraph_format.space_after = Pt(12)


def render_markdown(doc, chapter_path, config):
    text = chapter_path.read_text(encoding="utf-8")
    html = markdown.markdown(text, extensions=["tables", "fenced_code"])
    soup = BeautifulSoup(html, "html.parser")
    previous_was_image = False

    for node in soup.find_all(recursive=False):
        name = node.name
        if name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            add_heading(doc, int(name[1]), node.get_text(" ", strip=True), config)
        elif name == "p":
            image = node.find("img")
            if image is not None:
                if node.get_text(strip=True):
                    raise ValueError(
                        f"图片必须独占 Markdown 段落：{chapter_path}"
                    )
                add_image(doc, image.get("src", ""), chapter_path.parent, config)
                previous_was_image = True
                continue
            text_value = node.get_text()
            emphasis = node.find("em")
            if (
                previous_was_image
                and emphasis is not None
                and node.get_text(strip=True) == emphasis.get_text(strip=True)
            ):
                add_caption(doc, emphasis.get_text(" ", strip=True), config)
            else:
                if text_value.strip():
                    add_paragraph(doc, node, config)
        elif name in ("ul", "ol"):
            add_list(doc, node, config)
        elif name == "table":
            add_table(doc, node, config)
        elif name == "blockquote":
            add_blockquote(doc, node, config)
        elif name == "pre":
            add_code_block(doc, node.get_text(), config)
        elif name == "hr":
            paragraph = doc.add_paragraph("─" * 40)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            if node.get_text(strip=True):
                add_paragraph(doc, node, config)
        previous_was_image = False


def chapter_entries(config, base_dir):
    entries = []
    for raw in config["chapters"]:
        if isinstance(raw, str):
            raw = {"path": raw}
        if (
            not isinstance(raw, dict)
            or not isinstance(raw.get("path"), str)
            or not raw["path"].strip()
        ):
            raise ValueError("chapters 的每一项必须是路径字符串或包含 path 的对象")
        if "page_break_before" in raw and not isinstance(
            raw["page_break_before"], bool
        ):
            raise ValueError("page_break_before 必须是布尔值")
        path = (base_dir / raw["path"]).resolve()
        if not path.is_file():
            raise FileNotFoundError(f"章节不存在：{path}")
        entries.append((path, bool(raw.get("page_break_before", False))))
    return entries


def build(config_path, config):
    base_dir = config_path.parent
    chapters = chapter_entries(config, base_dir)
    output = (base_dir / config["output"]).resolve()
    if output.suffix.lower() != ".docx":
        raise ValueError("output 必须使用 .docx 扩展名")
    protected_paths = {config_path, *(chapter for chapter, _ in chapters)}
    if output in protected_paths:
        raise ValueError("output 不能覆盖配置文件或 Markdown 源文件")
    output.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    setup_document(doc, config)
    add_cover(doc, config)
    add_toc(doc, config)
    for index, (chapter, page_break_before) in enumerate(chapters):
        if page_break_before and index > 0:
            doc.add_page_break()
        render_markdown(doc, chapter, config)
    doc.save(output)
    return output


def main():
    config_path, config = load_config(parse_args().config)
    output = build(config_path, config)
    print(f"saved: {output}")


if __name__ == "__main__":
    main()
